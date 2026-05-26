import json
import io
import os
import re
import textwrap

import docx
import PyPDF2
from flask import Flask, abort, redirect, render_template, request, send_file, session
from werkzeug.utils import secure_filename
from werkzeug.security import check_password_hash, generate_password_hash

from ai import analyze_resume
from db import Base, engine, SessionLocal
import models


app = Flask(__name__)
app.secret_key = os.getenv("SECRET_KEY", "dev-only-change-me")

Base.metadata.create_all(bind=engine)


def _clean_list(value):
    if not isinstance(value, list):
        return []
    return [str(item).strip() for item in value if str(item).strip()]


def _fallback_careers(user_goal, skills):
    goal = user_goal.strip() or "Your target role"
    lowered_goal = goal.lower()
    recommendations = {
        "data": ["Data Analyst", "Business Intelligence Analyst", "Product Analyst"],
        "frontend": ["Frontend Developer", "UI Engineer", "Full Stack Developer"],
        "front end": ["Frontend Developer", "UI Engineer", "Full Stack Developer"],
        "backend": ["Backend Developer", "API Developer", "Cloud Application Developer"],
        "back end": ["Backend Developer", "API Developer", "Cloud Application Developer"],
        "full stack": ["Full Stack Developer", "Frontend Developer", "Backend Developer"],
        "python": ["Python Developer", "Backend Developer", "Automation Engineer"],
        "design": ["UX Designer", "UI Designer", "Product Designer"],
        "marketing": ["Digital Marketing Specialist", "Content Strategist", "SEO Specialist"],
    }
    titles = [goal, f"{goal} Specialist", f"Associate {goal}"]
    for keyword, choices in recommendations.items():
        if keyword in lowered_goal:
            titles = choices
            break

    skill_text = ", ".join(skills[:3]) or "your documented experience"
    return [
        {
            "title": title,
            "match_reason": f"Builds on {skill_text} while moving toward {goal}.",
            "skills_to_build": "Strengthen the missing skills highlighted in your roadmap."
        }
        for title in titles[:3]
    ]


def normalize_analysis(result, user_goal, resume_text):
    if not isinstance(result, dict) or result.get("error"):
        return result

    result["skills"] = _clean_list(result.get("skills"))
    result["missing_skills"] = _clean_list(result.get("missing_skills"))
    result["roadmap"] = _clean_list(result.get("roadmap"))
    result["interview_questions"] = _clean_list(result.get("interview_questions"))

    try:
        score = int(float(result.get("resume_score")))
    except (TypeError, ValueError):
        score = 48 + min(len(result["skills"]) * 7, 35) - min(
            len(result["missing_skills"]) * 4, 20
        )
    result["resume_score"] = max(0, min(100, score))

    if not str(result.get("score_summary") or "").strip():
        if result["resume_score"] >= 80:
            summary = "Strong role alignment. Refine impact statements to stand out."
        elif result["resume_score"] >= 60:
            summary = "Promising match. Closing the listed gaps can raise your fit."
        else:
            summary = "A developing match. Focus your resume on role-specific evidence."
        result["score_summary"] = summary

    careers = []
    raw_careers = result.get("career_options", [])
    if not isinstance(raw_careers, list):
        raw_careers = []
    for career in raw_careers:
        if isinstance(career, dict) and str(career.get("title") or "").strip():
            careers.append({
                "title": str(career["title"]).strip(),
                "match_reason": str(career.get("match_reason") or "").strip(),
                "skills_to_build": str(career.get("skills_to_build") or "").strip()
            })
        elif str(career).strip():
            careers.append({
                "title": str(career).strip(),
                "match_reason": "Matches skills identified in your resume.",
                "skills_to_build": "Follow your learning roadmap to strengthen the match."
            })
    result["career_options"] = careers[:3] or _fallback_careers(
        user_goal, result["skills"]
    )

    improved_resume = str(result.get("improved_resume") or "").strip()
    result["improved_resume"] = improved_resume or resume_text
    return result


def _pdf_escape(text):
    encoded = text.encode("cp1252", errors="replace").decode("latin-1")
    return encoded.replace("\\", "\\\\").replace("(", "\\(").replace(")", "\\)")


def build_resume_pdf(resume_text, role, heading_prefix="Optimized Resume"):
    heading = f"{heading_prefix} - {role or 'Career Profile'}"
    lines = [heading, "", "Generated from your analyzed resume", ""]
    for source_line in resume_text.splitlines() or [resume_text]:
        wrapped = textwrap.wrap(source_line, width=92) if source_line.strip() else [""]
        lines.extend(wrapped or [""])

    pages = [lines[index:index + 48] for index in range(0, len(lines), 48)] or [[]]
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica /Encoding /WinAnsiEncoding >>",
    ]
    page_ids = []
    for page_lines in pages:
        page_id = len(objects) + 1
        content_id = page_id + 1
        page_ids.append(page_id)
        commands = ["BT", "/F1 11 Tf", "52 790 Td", "15 TL"]
        for line in page_lines:
            commands.append(f"({_pdf_escape(line)}) Tj")
            commands.append("T*")
        commands.append("ET")
        stream = "\n".join(commands).encode("latin-1")
        objects.append(
            f"<< /Type /Page /Parent 2 0 R /MediaBox [0 0 612 842] "
            f"/Resources << /Font << /F1 3 0 R >> >> /Contents {content_id} 0 R >>".encode()
        )
        objects.append(
            f"<< /Length {len(stream)} >>\nstream\n".encode() + stream + b"\nendstream"
        )

    kids = " ".join(f"{page_id} 0 R" for page_id in page_ids)
    objects[1] = f"<< /Type /Pages /Count {len(page_ids)} /Kids [{kids}] >>".encode()

    pdf = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for number, obj in enumerate(objects, start=1):
        offsets.append(len(pdf))
        pdf.extend(f"{number} 0 obj\n".encode() + obj + b"\nendobj\n")
    xref_offset = len(pdf)
    pdf.extend(f"xref\n0 {len(objects) + 1}\n0000000000 65535 f \n".encode())
    for offset in offsets[1:]:
        pdf.extend(f"{offset:010d} 00000 n \n".encode())
    pdf.extend(
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF".encode()
    )
    return bytes(pdf)


COACH_STOP_WORDS = {
    "about", "after", "also", "ability", "able", "and", "are", "based",
    "candidate", "company", "experience", "from", "have", "including", "job",
    "must", "our", "role", "should", "strong", "team", "that", "the", "their",
    "this", "using", "with", "will", "work", "years", "your"
}


def _resume_from_upload(file):
    if not file or file.filename == "":
        return "", None

    filename = file.filename.lower()
    try:
        if filename.endswith(".pdf"):
            pdf_reader = PyPDF2.PdfReader(file)
            return "\n".join(page.extract_text() or "" for page in pdf_reader.pages), None
        if filename.endswith(".docx"):
            document = docx.Document(file)
            return "\n".join(paragraph.text for paragraph in document.paragraphs), None
    except Exception as exc:
        return "", f"Could not read this resume file: {str(exc)}"

    return "", "Please upload a PDF or DOCX resume file."


def _jd_keywords(job_description):
    terms = re.findall(r"[A-Za-z][A-Za-z+#.-]{2,}", job_description.lower())
    unique_terms = []
    for term in terms:
        cleaned = term.rstrip(".,")
        if cleaned not in COACH_STOP_WORDS and cleaned not in unique_terms:
            unique_terms.append(cleaned)
    return unique_terms[:22]


def _bounded_score(value):
    return max(0, min(100, int(round(value))))


def ideal_role_profile(target_role):
    role_lower = target_role.lower()
    profiles = {
        "backend": {
            "skills": ["Python", "REST API", "SQL", "Testing", "Docker", "Cloud", "Git"],
            "stages": [
                ("Foundation", 18, ["Python", "Git", "SQL"], "Master language fundamentals, databases, and version-controlled exercises."),
                ("Application", 26, ["REST APIs", "Authentication", "Testing"], "Build production-style APIs with documented tests and security basics."),
                ("Advanced", 35, ["Docker", "Cloud Deployment", "System Design"], "Deploy a scalable service and prepare architecture discussions."),
            ]
        },
        "frontend": {
            "skills": ["HTML", "CSS", "JavaScript", "React", "Accessibility", "Testing", "Performance"],
            "stages": [
                ("Foundation", 16, ["HTML", "CSS", "JavaScript"], "Build responsive, accessible interfaces from first principles."),
                ("Application", 24, ["React", "State Management", "Testing"], "Create interactive products with reliable component tests."),
                ("Advanced", 30, ["Performance", "Accessibility", "Deployment"], "Polish production quality, web vitals, and portfolio delivery."),
            ]
        },
        "data": {
            "skills": ["Excel", "SQL", "Python", "Statistics", "Visualization", "Power BI", "Storytelling"],
            "stages": [
                ("Foundation", 18, ["Excel", "SQL", "Statistics"], "Learn cleaning, querying, and interpreting core business data."),
                ("Application", 24, ["Python", "Visualization", "Power BI"], "Produce dashboards and data-backed recommendations."),
                ("Advanced", 32, ["Experimentation", "Forecasting", "Stakeholder Storytelling"], "Deliver an end-to-end analytical case study."),
            ]
        },
        "design": {
            "skills": ["Figma", "UX Research", "Wireframing", "Prototyping", "Accessibility", "Design Systems"],
            "stages": [
                ("Foundation", 14, ["Figma", "Visual Principles", "Wireframing"], "Practice clear layouts and essential interaction patterns."),
                ("Application", 24, ["UX Research", "Prototyping", "Usability Testing"], "Validate a product concept through user feedback."),
                ("Advanced", 30, ["Design Systems", "Accessibility", "Case Studies"], "Ship a professional portfolio narrative."),
            ]
        },
        "marketing": {
            "skills": ["SEO", "Content Strategy", "Analytics", "Campaigns", "Conversion", "Reporting"],
            "stages": [
                ("Foundation", 14, ["SEO", "Content", "Analytics"], "Learn acquisition channels and measurable objectives."),
                ("Application", 21, ["Campaign Planning", "Conversion", "Reporting"], "Create and evaluate a realistic campaign plan."),
                ("Advanced", 28, ["Growth Strategy", "Attribution", "Presentation"], "Build a strategic portfolio case study."),
            ]
        }
    }
    for keyword, profile in profiles.items():
        if keyword in role_lower or (keyword == "backend" and "api" in role_lower):
            return profile
    return {
        "skills": ["Communication", "Problem Solving", "Domain Knowledge", "Tools", "Projects", "Results"],
        "stages": [
            ("Foundation", 14, ["Core Concepts", "Essential Tools"], "Learn the fundamentals expected for this profession."),
            ("Application", 24, ["Practical Projects", "Role Skills"], "Build demonstrable evidence aligned to common responsibilities."),
            ("Advanced", 30, ["Specialization", "Interview Readiness"], "Polish your portfolio and prepare to explain your decisions."),
        ]
    }


def build_interview_questions(analysis, target_role, matched_terms, missing_terms):
    questions = _clean_list(analysis.get("interview_questions"))
    prompts = [
        f"Why are you a strong fit for a {target_role} position?",
        f"Walk me through a project where you applied {matched_terms[0] if matched_terms else 'your strongest relevant skill'}.",
        f"How are you developing capability in {missing_terms[0] if missing_terms else 'advanced role-specific work'}?",
        "Describe a difficult decision you made, the tradeoffs, and the measurable outcome.",
        f"What would your first 30 days look like in this {target_role} role?"
    ]
    for prompt in prompts:
        if prompt not in questions:
            questions.append(prompt)
    return questions[:8]


def build_role_roadmap(target_role):
    profile = ideal_role_profile(target_role)
    stages = []
    
    # Project templates based on stage
    project_by_stage = {
        "Foundation": {
            "name": "Foundational Learning Project",
            "description": f"Build a simple project applying the core {target_role} concepts you're learning. Focus on understanding fundamentals.",
            "duration": "7-10 days",
            "deliverable": "Documented learning project with clear explanations.",
            "focus_skills": "Core Concepts, Essential Tools"
        },
        "Application": {
            "name": "Practical Portfolio Project",
            "description": f"Create a more substantial project demonstrating {target_role} capabilities. Include a real problem, solution, and measurable outcome.",
            "duration": "14-21 days",
            "deliverable": "Deployed project with documentation and results.",
            "focus_skills": "Practical Projects, Role Skills"
        },
        "Advanced": {
            "name": "Advanced Role-Ready Project",
            "description": f"Build a production-quality project showcasing advanced {target_role} expertise. This becomes your portfolio centerpiece.",
            "duration": "21-30 days",
            "deliverable": "Polished portfolio case study with architecture and learnings.",
            "focus_skills": "Specialization, Interview Readiness"
        }
    }
    
    for level, days, skills, outcome in profile["stages"]:
        stage_data = {
            "level": level,
            "days": days,
            "skills": skills,
            "outcome": outcome
        }
        # Add project recommendation for this stage
        if level in project_by_stage:
            project_info = project_by_stage[level]
            stage_data["project"] = {
                "name": project_info["name"],
                "description": project_info["description"],
                "duration": project_info["duration"],
                "deliverable": project_info["deliverable"],
                "focus_skills": project_info["focus_skills"]
            }
        stages.append(stage_data)
    
    return {
        "target_role": target_role,
        "career_roadmap": stages,
        "roadmap_days": sum(stage["days"] for stage in stages),
        "core_skills": profile["skills"],
    }


def build_coach_assessment(analysis, resume_text, target_role):
    resume_lower = resume_text.lower()
    profile = ideal_role_profile(target_role)
    terms = profile["skills"]
    matched_terms = [term for term in terms if term.lower() in resume_lower]
    missing_terms = [term for term in terms if term.lower() not in resume_lower]
    keyword_score = _bounded_score(
        (len(matched_terms) / len(terms)) * 100 if terms else analysis.get("resume_score", 60)
    )

    expected_sections = {
        "Summary": ["summary", "profile", "objective"],
        "Experience": ["experience", "employment", "work history"],
        "Skills": ["skills", "technical skills", "technologies"],
        "Projects": ["projects", "portfolio"],
        "Education": ["education", "degree", "university"],
    }
    found_sections = [
        name for name, markers in expected_sections.items()
        if any(marker in resume_lower for marker in markers)
    ]
    structure_score = _bounded_score(28 + len(found_sections) * 14)
    quantified_results = len(re.findall(r"\b\d+(?:\.\d+)?%|\b\d+\+?\s*(?:users|clients|projects|hours|days|requests|sales)", resume_lower))
    power_words = [
        "built", "delivered", "designed", "developed", "optimized", "implemented",
        "increased", "reduced", "launched", "led", "automated", "architected"
    ]
    verb_count = sum(resume_lower.count(word) for word in power_words)
    impact_score = _bounded_score(34 + quantified_results * 16 + verb_count * 5)
    skills = _clean_list(analysis.get("skills"))
    missing_skills = _clean_list(analysis.get("missing_skills"))
    skill_score = _bounded_score(48 + len(skills) * 9 - len(missing_skills) * 6)
    word_count = len(resume_text.split())
    readability_score = _bounded_score(84 if 180 <= word_count <= 850 else 64)

    charts = [
        {"title": "Ideal role skill match", "score": keyword_score, "note": f"{len(matched_terms)}/{len(terms)} ideal skills visible"},
        {"title": "Experience impact", "score": impact_score, "note": f"{quantified_results} quantified outcomes found"},
        {"title": "ATS structure", "score": structure_score, "note": f"{len(found_sections)}/5 core sections found"},
        {"title": "Skills coverage", "score": skill_score, "note": f"{len(skills)} aligned skills found"},
        {"title": "Readability", "score": readability_score, "note": f"{word_count} words reviewed"},
    ]
    overall_score = _bounded_score(
        keyword_score * 0.30
        + impact_score * 0.22
        + structure_score * 0.18
        + skill_score * 0.20
        + readability_score * 0.10
    )
    projected_score = _bounded_score(
        overall_score + min(25, len(missing_terms[:6]) * 3 + (100 - impact_score) * 0.12)
    )

    feedback_templates = {
        "Ideal role skill match": (
            "Ideal-role alignment",
            "Your updated resume does not yet show enough of the expected skills for this job.",
            "Add truthful evidence for the missing ideal-role skills through bullets or projects."
        ),
        "Experience impact": (
            "Experience bullets",
            "Impact is difficult to measure from the current achievement statements.",
            "Add scope, result, speed, quality, or scale to bullets using numbers only when accurate."
        ),
        "ATS structure": (
            "Resume structure",
            "Important standard sections may be missing or difficult for an ATS to recognize.",
            "Use clear headings such as Summary, Skills, Experience, Projects, and Education."
        ),
        "Skills coverage": (
            "Technical skill coverage",
            "Your visible skills do not fully support this target role yet.",
            "Prioritize the missing skills and prove them through one focused project."
        ),
        "Readability": (
            "Clarity and length",
            "The resume length may make the key story harder to scan.",
            "Keep bullets concise and front-load the skills and achievements most relevant to this role."
        ),
    }
    weak_sections = []
    for item in sorted(charts, key=lambda chart: chart["score"])[:3]:
        title, weakness, improvement = feedback_templates[item["title"]]
        weak_sections.append({
            "title": title,
            "score": item["score"],
            "weakness": weakness,
            "improvement": improvement
        })

    role_lower = target_role.lower()
    verb_sets = [
        {
            "trigger": ["develop", "software", "backend", "frontend", "api", "engineer"],
            "verbs": ["Architected", "Engineered", "Optimized", "Deployed"],
            "example": "Engineered [feature/system] using [technology] to improve [measurable outcome]."
        },
        {
            "trigger": ["data", "analyst", "analytics", "sql", "dashboard"],
            "verbs": ["Analyzed", "Modeled", "Visualized", "Forecasted"],
            "example": "Analyzed [dataset] to uncover [insight] and support [business decision]."
        },
        {
            "trigger": ["market", "content", "campaign", "growth", "seo"],
            "verbs": ["Launched", "Accelerated", "Converted", "Strategized"],
            "example": "Launched [campaign] that improved [metric] through [strategy]."
        },
    ]
    action_verb_groups = [
        group for group in verb_sets if any(trigger in role_lower for trigger in group["trigger"])
    ] or [{
        "verbs": ["Delivered", "Led", "Improved", "Streamlined"],
        "example": "Delivered [initiative] that achieved [accurate measurable result]."
    }]

    focus_skills = (missing_skills + missing_terms)[:5]
    if not focus_skills:
        focus_skills = ["advanced role-specific execution", "portfolio evidence"]
    
    projects_data = [
        {
            "level": "Foundational",
            "name": f"Starter {target_role} Project",
            "focus_skills": focus_skills[:2],
            "duration": "7-10 days",
            "description": f"Build a focused, complete project that demonstrates core {target_role} skills. Start simple, document clearly, and prepare it for portfolio sharing.",
            "required_skills": focus_skills[:2] if focus_skills else ["core fundamentals"],
            "steps": [
                f"Identify a real problem or task relevant to {target_role}.",
                "Define a clear scope and success criteria.",
                f"Use essential {target_role} tools and technologies you're learning.",
                "Document what you built, why you chose it, and what you learned."
            ]
        },
        {
            "level": "Intermediate",
            "name": f"Production-Ready {target_role} Solution",
            "focus_skills": focus_skills[1:3] if len(focus_skills) > 1 else focus_skills,
            "duration": "10-14 days",
            "description": f"Create a polished, deployable {target_role} project that mirrors real job responsibilities. Include testing, documentation, and measurable outcomes.",
            "required_skills": (focus_skills[1:3] if len(focus_skills) > 1 else focus_skills) or ["intermediate capabilities"],
            "steps": [
                f"Design a realistic {target_role} workflow or application.",
                "Implement with best practices: clean code, testing, version control.",
                f"Deploy or deliver the complete {target_role} solution.",
                "Write a case study showing problem, approach, and results."
            ]
        },
        {
            "level": "Advanced",
            "name": f"Advanced {target_role} Portfolio Showcase",
            "focus_skills": focus_skills[2:5] if len(focus_skills) > 2 else focus_skills,
            "duration": "14-18 days",
            "description": f"Deliver a sophisticated {target_role} project that demonstrates system design, optimization, and leadership. This becomes your flagship portfolio piece.",
            "required_skills": (focus_skills[2:5] if len(focus_skills) > 2 else focus_skills) or ["advanced expertise"],
            "steps": [
                f"Design a scalable, well-architected {target_role} system.",
                "Implement advanced features, performance optimization, and monitoring.",
                "Present quantified impact: speed improvements, user reach, cost savings.",
                "Create a professional writeup and demo that showcases decision-making."
            ]
        }
    ]
    
    recommended_projects = [
        {
            "name": project["name"],
            "level": project["level"],
            "focus": ", ".join(project["focus_skills"]),
            "duration": project["duration"],
            "description": project["description"],
            "required_skills": project["required_skills"],
            "steps": project["steps"]
        }
        for project in projects_data
    ]
    strong_parts = []
    if matched_terms:
        strong_parts.append({
            "title": "Role-aligned skills",
            "detail": "Your resume already demonstrates " + ", ".join(matched_terms[:4]) + "."
        })
    if quantified_results:
        strong_parts.append({
            "title": "Evidence of impact",
            "detail": f"{quantified_results} quantified outcome(s) help your achievements stand out."
        })
    if len(found_sections) >= 4:
        strong_parts.append({
            "title": "ATS-friendly structure",
            "detail": "Clear core resume sections improve readability and parsing."
        })
    if not strong_parts:
        strong_parts.append({
            "title": "Professional foundation",
            "detail": "Your content provides a base the coach can sharpen for this target role."
        })

    coached_resume = str(analysis.get("improved_resume") or "").strip() or resume_text
    interview_questions = build_interview_questions(
        analysis, target_role, matched_terms, missing_terms
    )

    coach_result = dict(analysis)
    coach_result.update({
        "target_role": target_role,
        "compatibility_score": overall_score,
        "projected_score": projected_score,
        "charts": charts,
        "strong_parts": strong_parts,
        "weak_sections": weak_sections,
        "matched_keywords": matched_terms[:10],
        "missing_keywords": missing_terms[:10],
        "action_verb_groups": action_verb_groups,
        "recommended_projects": recommended_projects,
        "interview_questions": interview_questions,
        "coach_resume": coached_resume,
    })
    return coach_result


def latest_coach_result(email):
    db = SessionLocal()
    try:
        user = db.query(models.User).filter_by(email=email).first()
        if not user:
            return None
        report = db.query(models.CoachReport).filter_by(
            user_id=user.id
        ).order_by(models.CoachReport.id.desc()).first()
        if not report:
            return None
        try:
            result = json.loads(report.results or "{}")
        except json.JSONDecodeError:
            result = {}
        result["coach_id"] = report.id
        return result
    finally:
        db.close()


# HOME ROUTE
@app.route("/")
def home():

    if "user" in session:
        return redirect("/dashboard")

    return redirect("/login")


# SIGNUP ROUTE
@app.route("/signup", methods=["GET", "POST"])
def signup():
    db = SessionLocal()
    try:
        if request.method == "POST":
            email = (request.form.get("email") or "").strip().lower()
            password = request.form.get("password") or ""

            if len(password) < 6:
                return render_template(
                    "signup.html",
                    error="Password must be at least 6 characters."
                )

            existing_user = db.query(models.User).filter_by(email=email).first()
            if existing_user:
                return render_template("signup.html", error="User already exists.")

            user = models.User(
                email=email,
                password=generate_password_hash(password)
            )
            db.add(user)
            db.commit()
            return redirect("/login")

        return render_template("signup.html")
    finally:
        db.close()


# LOGIN ROUTE
@app.route("/login", methods=["GET", "POST"])
def login():
    db = SessionLocal()
    try:
        if request.method == "POST":
            email = (request.form.get("email") or "").strip().lower()
            password = request.form.get("password") or ""
            user = db.query(models.User).filter_by(email=email).first()

            valid_password = False
            if user:
                if user.password == password:
                    valid_password = True
                else:
                    try:
                        valid_password = check_password_hash(user.password, password)
                    except ValueError:
                        valid_password = False

            if valid_password:
                if user.password == password:
                    user.password = generate_password_hash(password)
                    db.commit()
                session["user"] = user.email
                return redirect("/dashboard")

            return render_template("login.html", error="Invalid credentials.")

        return render_template("login.html")
    finally:
        db.close()


# FORGOT_PASSWORD ROUTE
@app.route("/forgot_password", methods=["GET", "POST"])
def forgot_password():
    if request.method == "POST":
        return render_template(
            "forgot_password.html",
            message="Password reset email delivery is not configured yet."
        )
    return render_template("forgot_password.html")


# DASHBOARD ROUTE

@app.route("/dashboard", methods=["GET", "POST"])
def dashboard():

    if "user" not in session:
        return redirect("/login")

    result = None
    report_id = None
    resume_text = ""
    user_goal = ""

    if request.method == "POST":

        user_goal = (request.form.get("role") or "").strip()
        resume_text = (request.form.get("resume") or "").strip()

        file = request.files.get("file")

        # ==========================================
        # FILE HANDLING
        # ==========================================

        if file and file.filename != "":

            filename = file.filename.lower()

            # PDF FILE
            if filename.endswith(".pdf"):

                try:

                    pdf_reader = PyPDF2.PdfReader(file)

                    text = ""

                    for page in pdf_reader.pages:
                        text += page.extract_text() or ""

                    resume_text = text

                except Exception as e:

                    result = {
                        "error": f"PDF Error: {str(e)}"
                    }

            # DOCX FILE
            elif filename.endswith(".docx"):

                try:

                    doc = docx.Document(file)

                    text = ""

                    for para in doc.paragraphs:
                        text += para.text + "\n"

                    resume_text = text

                except Exception as e:

                    result = {
                        "error": f"DOCX Error: {str(e)}"
                    }
            else:
                result = {"error": "Please upload a PDF or DOCX file."}

        # ==========================================
        # AI ANALYSIS
        # ==========================================

        if result is None and resume_text and user_goal:

            try:

                result = analyze_resume(
                    resume_text,
                    user_goal
                )
                result = normalize_analysis(result, user_goal, resume_text)

                # ==========================================
                # SAVE TO DATABASE
                # ==========================================

                try:

                    db = SessionLocal()
                    try:
                        user = db.query(models.User).filter_by(
                            email=session["user"]
                        ).first()

                        if user and not result.get("error"):
                            report = models.Report(
                                user_id=user.id,
                                title=user_goal,
                                resume_text=resume_text,
                                results=json.dumps(result)
                            )
                            db.add(report)
                            db.commit()
                            report_id = report.id
                    finally:
                        db.close()

                except Exception as e:
                    result["error"] = f"Database save error: {str(e)}"

            except Exception as e:

                result = {
                    "error": f"AI Error: {str(e)}"
                }
        elif result is None:
            result = {"error": "Provide both a resume and a target role."}

    return render_template(
        "dashboard.html",
        user=session["user"],
        result=result,
        report_id=report_id
    )


@app.route("/coach", methods=["GET", "POST"])
def coach():
    if "user" not in session:
        return redirect("/login")

    result = None
    target_role = ""
    resume_text = ""

    if request.method == "POST":
        target_role = (request.form.get("role") or "").strip()
        resume_text = (request.form.get("resume") or "").strip()
        uploaded_text, upload_error = _resume_from_upload(request.files.get("file"))
        if upload_error:
            result = {"error": upload_error}
        elif uploaded_text:
            resume_text = uploaded_text

        if result is None and (not target_role or not resume_text):
            result = {
                "error": "Provide a target job and your updated resume to begin coaching."
            }

        if result is None:
            try:
                analysis = normalize_analysis(
                    analyze_resume(resume_text, target_role),
                    target_role,
                    resume_text
                )
                if analysis.get("error"):
                    result = analysis
                else:
                    result = build_coach_assessment(
                        analysis, resume_text, target_role
                    )
                    db = SessionLocal()
                    try:
                        user = db.query(models.User).filter_by(email=session["user"]).first()
                        if user:
                            report = models.CoachReport(
                                user_id=user.id,
                                target_role=target_role,
                                job_description="",
                                resume_text=resume_text,
                                results=json.dumps(result)
                            )
                            db.add(report)
                            db.commit()
                            result["coach_id"] = report.id
                    finally:
                        db.close()
            except Exception as exc:
                result = {"error": f"Coach analysis error: {str(exc)}"}

    return render_template(
        "coach.html",
        result=result,
        latest=None if result else latest_coach_result(session["user"]),
        target_role=target_role,
        resume_text=resume_text
    )


@app.route("/action-verbs")
def action_verbs():
    if "user" not in session:
        return redirect("/login")
    return render_template(
        "action_verbs.html",
        result=latest_coach_result(session["user"])
    )


@app.route("/projects")
def projects():
    if "user" not in session:
        return redirect("/login")
    return render_template(
        "projects.html",
        result=latest_coach_result(session["user"])
    )


@app.route("/roadmap", methods=["GET", "POST"])
def roadmap():
    if "user" not in session:
        return redirect("/login")
    target_role = ""
    result = None
    if request.method == "POST":
        target_role = (request.form.get("role") or "").strip()
        if target_role:
            result = build_role_roadmap(target_role)
        else:
            result = {"error": "Enter your desired job to create a curated roadmap."}
    return render_template(
        "roadmap.html",
        result=result,
        target_role=target_role
    )


def _user_coach_report(coach_id):
    db = SessionLocal()
    user = db.query(models.User).filter_by(email=session["user"]).first()
    report = None
    if user:
        report = db.query(models.CoachReport).filter_by(
            id=coach_id,
            user_id=user.id
        ).first()
    return db, report


@app.route("/coach/<int:coach_id>/resume.pdf")
def download_coach_resume(coach_id):
    if "user" not in session:
        return redirect("/login")
    db, report = _user_coach_report(coach_id)
    try:
        if not report:
            abort(404)
        result = json.loads(report.results or "{}")
        resume_text = result.get("coach_resume") or result.get("improved_resume") or report.resume_text
        filename = secure_filename(f"{report.target_role or 'coach'}_coach_resume.pdf")
        return send_file(
            io.BytesIO(build_resume_pdf(resume_text, report.target_role, "Resume Given by AI Coach")),
            mimetype="application/pdf",
            as_attachment=True,
            download_name=filename or "ai_coach_resume.pdf"
        )
    finally:
        db.close()


@app.route("/coach/<int:coach_id>/interview-questions.pdf")
def download_interview_questions(coach_id):
    if "user" not in session:
        return redirect("/login")
    db, report = _user_coach_report(coach_id)
    try:
        if not report:
            abort(404)
        result = json.loads(report.results or "{}")
        questions = _clean_list(result.get("interview_questions"))
        content = "\n\n".join(
            f"{index}. {question}\nPreparation notes: "
            for index, question in enumerate(questions, start=1)
        )
        filename = secure_filename(f"{report.target_role or 'role'}_interview_questions.pdf")
        return send_file(
            io.BytesIO(build_resume_pdf(content, report.target_role, "Important Interview Questions")),
            mimetype="application/pdf",
            as_attachment=True,
            download_name=filename or "interview_questions.pdf"
        )
    finally:
        db.close()


@app.route("/report/<int:report_id>/download")
def download_report(report_id):
    if "user" not in session:
        return redirect("/login")

    db = SessionLocal()
    try:
        user = db.query(models.User).filter_by(email=session["user"]).first()
        if not user:
            session.pop("user", None)
            return redirect("/login")

        report = db.query(models.Report).filter_by(
            id=report_id,
            user_id=user.id
        ).first()
        if not report:
            abort(404)

        try:
            parsed_result = json.loads(report.results or "{}")
        except json.JSONDecodeError:
            parsed_result = {}
        resume_text = parsed_result.get("improved_resume") or report.resume_text or ""
        filename = secure_filename(f"{report.title or 'optimized'}_resume.pdf")
        return send_file(
            io.BytesIO(build_resume_pdf(resume_text, report.title)),
            mimetype="application/pdf",
            as_attachment=True,
            download_name=filename or "optimized_resume.pdf"
        )
    finally:
        db.close()


@app.route("/history")
def history():

    # check login
    if "user" not in session:
        return redirect("/login")

    db = SessionLocal()
    try:
        user = db.query(models.User).filter_by(
            email=session["user"]
        ).first()

        if not user:
            session.pop("user", None)
            return redirect("/login")

        reports = db.query(models.Report).filter_by(
            user_id=user.id
        ).order_by(models.Report.id.desc()).all()

        parsed_reports = []
        for r in reports:
            try:
                parsed_result = json.loads(r.results)
            except (TypeError, json.JSONDecodeError):
                parsed_result = {}

            parsed_reports.append({
                "id": r.id,
                "title": r.title,
                "resume": r.resume_text,
                "result": parsed_result
            })

        return render_template(
            "history.html",
            reports=parsed_reports,
            user=session["user"]
        )
    finally:
        db.close()


# LOGOUT ROUTE
@app.route("/logout")
def logout():

    session.pop("user", None)

    return redirect("/login")


if __name__ == "__main__":
    host = os.getenv("HOST", "0.0.0.0")
    port = int(os.getenv("PORT", "5000"))
    debug = os.getenv("FLASK_DEBUG", "").lower() in {"1", "true", "yes"}
    app.run(host=host, port=port, debug=debug, use_reloader=debug)
